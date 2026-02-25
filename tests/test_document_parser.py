"""Tests for real document parsing module."""

from __future__ import annotations

import io

import pytest

from app.document_parser import (
    DocumentChunk,
    PageBlock,
    chunk_text_blocks,
    parse_file_bytes,
    parse_pdf_bytes,
    parse_plain_text_bytes,
)


class TestChunkTextBlocks:
    def test_empty_blocks(self):
        result = chunk_text_blocks([])
        assert result == []

    def test_single_short_block(self):
        blocks = [PageBlock(page=1, text="Hello world. " * 10, bbox=[0.0, 0.0, 0.5, 0.5])]
        result = chunk_text_blocks(blocks, chunk_size=500, chunk_overlap=50)
        assert len(result) == 1
        assert result[0].page == 1
        assert "Hello world" in result[0].text

    def test_multiple_blocks_produce_multiple_chunks(self):
        blocks = []
        for i in range(20):
            blocks.append(
                PageBlock(
                    page=(i // 5) + 1,
                    text=f"Paragraph {i}. " + "This is content that fills up the chunk. " * 5,
                    bbox=[0.0, 0.1 * (i % 5), 1.0, 0.1 * (i % 5) + 0.1],
                )
            )
        result = chunk_text_blocks(blocks, chunk_size=300, chunk_overlap=50)
        assert len(result) > 1
        for chunk in result:
            assert chunk.page >= 1
            assert len(chunk.text) >= 50
            assert chunk.chunk_id.startswith("ck_")

    def test_chunk_overlap(self):
        long_text = "Word " * 200
        blocks = [PageBlock(page=1, text=long_text, bbox=[0.0, 0.0, 1.0, 1.0])]
        result = chunk_text_blocks(blocks, chunk_size=100, chunk_overlap=30)
        if len(result) >= 2:
            end_of_first = result[0].text[-30:]
            assert end_of_first in result[1].text

    def test_preserves_page_numbers(self):
        blocks = [
            PageBlock(page=1, text="Page one content. " * 30, bbox=[0.0, 0.0, 1.0, 0.5]),
            PageBlock(page=2, text="Page two content. " * 30, bbox=[0.0, 0.0, 1.0, 0.5]),
            PageBlock(page=3, text="Page three content. " * 30, bbox=[0.0, 0.0, 1.0, 0.5]),
        ]
        result = chunk_text_blocks(blocks, chunk_size=200, chunk_overlap=30)
        pages_seen = {chunk.page for chunk in result}
        assert len(pages_seen) >= 2


class TestParsePdfBytes:
    def _make_simple_pdf(self, text: str = "Hello PDF World") -> bytes:
        """Create a minimal PDF with text using pymupdf."""
        import pymupdf

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), text, fontsize=12)
        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        return buf.getvalue()

    def test_parse_simple_pdf(self):
        long_text = "This is a test document for bid evaluation. The supplier qualification review has been completed."
        pdf_bytes = self._make_simple_pdf(long_text)
        chunks = parse_pdf_bytes(pdf_bytes)
        assert len(chunks) >= 1
        full_text = " ".join(c.text for c in chunks)
        assert "test document" in full_text

    def test_parse_multi_page_pdf(self):
        import pymupdf

        doc = pymupdf.open()
        for i in range(3):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i + 1} content here. " * 20, fontsize=11)
        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        chunks = parse_pdf_bytes(buf.getvalue())
        assert len(chunks) >= 1
        pages_seen = {c.page for c in chunks}
        assert len(pages_seen) >= 1

    def test_parse_empty_pdf(self):
        import pymupdf

        doc = pymupdf.open()
        doc.new_page()
        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        chunks = parse_pdf_bytes(buf.getvalue())
        assert chunks == []

    def test_parse_corrupt_pdf(self):
        from app.errors import ApiError

        with pytest.raises(ApiError) as exc_info:
            parse_pdf_bytes(b"not a pdf file")
        assert exc_info.value.code == "DOC_PARSE_PDF_CORRUPT"


class TestParsePlainText:
    def test_parse_simple_text(self):
        text = "第一段内容。" * 20 + "\n\n" + "第二段内容。" * 20
        chunks = parse_plain_text_bytes(text.encode("utf-8"))
        assert len(chunks) >= 1
        full_text = " ".join(c.text for c in chunks)
        assert "第一段" in full_text

    def test_parse_empty_text(self):
        chunks = parse_plain_text_bytes(b"")
        assert chunks == []

    def test_parse_gb18030_text(self):
        text = "这是一份GB18030编码的文档内容。" * 20
        chunks = parse_plain_text_bytes(text.encode("gb18030"))
        assert len(chunks) >= 1
        assert "GB18030" in " ".join(c.text for c in chunks)


class TestParseFileBytes:
    def _make_simple_pdf(self) -> bytes:
        import pymupdf

        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 72), "供应商资质认证报告\n\n" + "详细内容 " * 50, fontsize=11)
        buf = io.BytesIO()
        doc.save(buf)
        doc.close()
        return buf.getvalue()

    def test_parse_pdf_file(self):
        pdf_bytes = self._make_simple_pdf()
        result = parse_file_bytes(
            pdf_bytes,
            filename="report.pdf",
            document_id="doc_test001",
        )
        assert len(result) >= 1
        chunk = result[0]
        assert chunk["document_id"] == "doc_test001"
        assert chunk["parser"] == "local"
        assert chunk["content_source"] == "pdf_local"
        assert "chunk_id" in chunk
        assert "chunk_hash" in chunk
        assert isinstance(chunk["positions"], list)
        assert isinstance(chunk["heading_path"], list)

    def test_parse_txt_file(self):
        text = "评标标准第一条\n\n" + "详细描述内容。" * 30
        result = parse_file_bytes(
            text.encode("utf-8"),
            filename="criteria.txt",
            document_id="doc_test002",
        )
        assert len(result) >= 1
        assert result[0]["content_source"] == "text_local"

    def test_chunk_schema_completeness(self):
        text = "内容 " * 100
        result = parse_file_bytes(
            text.encode("utf-8"),
            filename="doc.txt",
            document_id="doc_schema_test",
        )
        required_keys = {
            "chunk_id",
            "document_id",
            "pages",
            "positions",
            "section",
            "heading_path",
            "chunk_type",
            "parser",
            "parser_version",
            "content_source",
            "text",
            "chunk_hash",
        }
        for chunk in result:
            assert required_keys.issubset(chunk.keys()), f"Missing keys: {required_keys - chunk.keys()}"


class TestDocxParsing:
    def _make_simple_docx(self) -> bytes:
        try:
            import docx
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = docx.Document()
        doc.add_heading("评标报告", level=1)
        doc.add_paragraph("供应商 A 的资质审查结论如下。" * 10)
        doc.add_heading("技术评分", level=2)
        doc.add_paragraph("技术方案评审得分为 85 分。" * 10)
        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def test_parse_docx(self):
        from app.document_parser import parse_docx_bytes

        docx_bytes = self._make_simple_docx()
        chunks = parse_docx_bytes(docx_bytes)
        assert len(chunks) >= 1
        full_text = " ".join(c.text for c in chunks)
        assert "评标报告" in full_text or "资质" in full_text

    def test_parse_docx_file_bytes(self):
        docx_bytes = self._make_simple_docx()
        result = parse_file_bytes(
            docx_bytes,
            filename="report.docx",
            document_id="doc_docx_001",
        )
        assert len(result) >= 1
        assert result[0]["content_source"] == "docx_local"
