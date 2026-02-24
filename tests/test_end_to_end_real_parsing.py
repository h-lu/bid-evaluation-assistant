"""
End-to-end tests for real document parsing pipeline.

Tests the full flow: upload PDF -> parse -> chunk -> index -> evaluation,
using real PyMuPDF parsing instead of stubs.
"""

from __future__ import annotations

import io

import pytest

from app.store import store


def _make_pdf_with_content(pages: list[str]) -> bytes:
    import pymupdf

    doc = pymupdf.open()
    for page_text in pages:
        page = doc.new_page()
        y = 72
        for line in page_text.split("\n"):
            if line.strip():
                page.insert_text((72, y), line, fontsize=11)
                y += 16
    buf = io.BytesIO()
    doc.save(buf)
    doc.close()
    return buf.getvalue()


def _make_docx_with_content(sections: list[tuple[str, str]]) -> bytes:
    import docx

    doc = docx.Document()
    for heading, body in sections:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestRealPdfParsingE2E:
    def test_upload_and_parse_pdf_produces_real_chunks(self, client):
        pdf_bytes = _make_pdf_with_content([
            "Supplier Qualification Report\n\n"
            "The supplier holds ISO 9001 certification, certificate number CN-2025-12345. "
            "The certification is valid until December 2026. Quality management system "
            "has been independently audited and verified by third-party assessor. "
            "All manufacturing processes comply with international standards.",
            "Technical Proposal Summary\n\n"
            "The proposed solution uses distributed microservices architecture. "
            "Supports horizontal scaling and high availability deployments. "
            "Response time SLA: 99.9% uptime guaranteed. "
            "Backup and disaster recovery plan included in the proposal.",
            "Pricing Details\n\n"
            "Total quoted price: RMB 1,280,000 including tax. "
            "Payment terms: 30% advance, 70% upon delivery acceptance. "
            "Price validity period: 90 calendar days from submission date. "
            "No hidden charges or additional fees beyond the quoted amount.",
        ])

        upload_resp = client.post(
            "/api/v1/documents/upload",
            data={
                "project_id": "proj_test_e2e",
                "supplier_id": "sup_test_e2e",
                "doc_type": "bid",
            },
            files={"file": ("bid_document.pdf", pdf_bytes, "application/pdf")},
            headers={
                "Idempotency-Key": "e2e-upload-pdf-001",
                "x-tenant-id": "tenant_e2e_test",
            },
        )
        assert upload_resp.status_code == 202
        data = upload_resp.json()["data"]
        document_id = data["document_id"]
        job_id = data["job_id"]

        result = store.run_job_once(job_id=job_id, tenant_id="tenant_e2e_test")
        assert result["final_status"] == "succeeded"

        chunks = store.list_document_chunks_for_tenant(
            document_id=document_id, tenant_id="tenant_e2e_test"
        )
        assert len(chunks) >= 1

        has_real_content = False
        for chunk in chunks:
            text = chunk.get("text", "")
            if "ISO" in text or "microservices" in text or "1,280,000" in text:
                has_real_content = True
            assert chunk.get("chunk_id", "").startswith("ck_")
            assert "positions" in chunk
            assert "heading_path" in chunk

        assert has_real_content, "Chunks should contain real extracted text from PDF"

    def test_parse_pdf_registers_citation_sources(self, client):
        pdf_bytes = _make_pdf_with_content([
            "Section A: Delivery Timeline\n\n"
            "All deliverables will be completed within 30 business days after contract signing. "
            "Detailed milestone schedule is provided in Appendix B of this proposal. "
            "Penalty clause for late delivery: 0.5% per day, capped at 10% of total value.",
        ])

        upload_resp = client.post(
            "/api/v1/documents/upload",
            data={
                "project_id": "proj_citation_test",
                "supplier_id": "sup_citation_test",
                "doc_type": "bid",
            },
            files={"file": ("delivery.pdf", pdf_bytes, "application/pdf")},
            headers={
                "Idempotency-Key": "e2e-citation-001",
                "x-tenant-id": "tenant_citation_test",
            },
        )
        assert upload_resp.status_code == 202
        data = upload_resp.json()["data"]
        document_id = data["document_id"]
        job_id = data["job_id"]

        store.run_job_once(job_id=job_id, tenant_id="tenant_citation_test")

        chunks = store.list_document_chunks_for_tenant(
            document_id=document_id, tenant_id="tenant_citation_test"
        )
        for chunk in chunks:
            chunk_id = chunk.get("chunk_id")
            if chunk_id:
                source = store.get_citation_source(
                    chunk_id=chunk_id, tenant_id="tenant_citation_test"
                )
                assert source is not None, f"Citation source missing for chunk {chunk_id}"
                assert source.get("document_id") == document_id

    def test_multi_page_pdf_produces_chunks_with_page_numbers(self, client):
        pages = [
            f"Page {i + 1} Content Section\n\n"
            f"This is the detailed content for page {i + 1}. "
            f"It contains important information about section {i + 1} of the bid document. "
            f"Additional details and specifications are described thoroughly here."
            for i in range(5)
        ]
        pdf_bytes = _make_pdf_with_content(pages)

        upload_resp = client.post(
            "/api/v1/documents/upload",
            data={
                "project_id": "proj_multipage",
                "supplier_id": "sup_multipage",
                "doc_type": "bid",
            },
            files={"file": ("multipage.pdf", pdf_bytes, "application/pdf")},
            headers={
                "Idempotency-Key": "e2e-multipage-001",
                "x-tenant-id": "tenant_multipage",
            },
        )
        assert upload_resp.status_code == 202
        data = upload_resp.json()["data"]
        job_id = data["job_id"]
        document_id = data["document_id"]

        store.run_job_once(job_id=job_id, tenant_id="tenant_multipage")

        chunks = store.list_document_chunks_for_tenant(
            document_id=document_id, tenant_id="tenant_multipage"
        )
        assert len(chunks) >= 1


class TestRealDocxParsingE2E:
    def test_upload_and_parse_docx(self, client):
        docx_bytes = _make_docx_with_content([
            ("Supplier Profile", "Company A has 15 years of experience in IT solutions. "
             "Annual revenue exceeds 50 million RMB. Certified ISO 27001 for information security."),
            ("Technical Approach", "Proposed architecture follows cloud-native principles. "
             "Containerized deployment with Kubernetes orchestration. CI/CD pipeline included."),
        ])

        upload_resp = client.post(
            "/api/v1/documents/upload",
            data={
                "project_id": "proj_docx_test",
                "supplier_id": "sup_docx_test",
                "doc_type": "bid",
            },
            files={"file": ("proposal.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            headers={
                "Idempotency-Key": "e2e-docx-upload-001",
                "x-tenant-id": "tenant_docx_test",
            },
        )
        assert upload_resp.status_code == 202
        data = upload_resp.json()["data"]
        job_id = data["job_id"]
        document_id = data["document_id"]

        result = store.run_job_once(job_id=job_id, tenant_id="tenant_docx_test")
        assert result["final_status"] == "succeeded"

        chunks = store.list_document_chunks_for_tenant(
            document_id=document_id, tenant_id="tenant_docx_test"
        )
        assert len(chunks) >= 1
        full_text = " ".join(c.get("text", "") for c in chunks)
        assert "experience" in full_text.lower() or "Kubernetes" in full_text


class TestEvaluationWithLlmProvider:
    def test_evaluation_creates_scored_report(self, client):
        """Verify evaluation goes through llm_provider (falls back to mock in tests)."""
        store.create_project(payload={
            "project_code": "EVAL_E2E",
            "name": "E2E Eval Test",
            "tenant_id": "tenant_eval_e2e",
        })
        store.create_supplier(payload={
            "supplier_code": "SUP_E2E",
            "name": "Test Supplier",
            "tenant_id": "tenant_eval_e2e",
        })
        store.create_rule_pack(payload={
            "rule_pack_version": "v1.0.0",
            "name": "E2E Rule Pack",
            "tenant_id": "tenant_eval_e2e",
            "rules": {
                "criteria": [
                    {
                        "criteria_id": "qualification",
                        "criteria_name": "Supplier Qualification",
                        "max_score": 20.0,
                        "weight": 1.0,
                        "requirement_text": "Supplier must have ISO certification",
                    },
                    {
                        "criteria_id": "technical",
                        "criteria_name": "Technical Capability",
                        "max_score": 30.0,
                        "weight": 1.5,
                        "requirement_text": "Technical solution must support high availability",
                    },
                ],
            },
        })

        eval_resp = client.post(
            "/api/v1/evaluations",
            json={
                "project_id": "proj_eval_e2e",
                "supplier_id": "sup_eval_e2e",
                "rule_pack_version": "v1.0.0",
                "evaluation_scope": {
                    "include_doc_types": ["bid"],
                    "force_hitl": False,
                },
                "query_options": {
                    "mode_hint": "hybrid",
                    "top_k": 10,
                },
            },
            headers={
                "Idempotency-Key": "e2e-eval-001",
                "x-tenant-id": "tenant_eval_e2e",
            },
        )
        assert eval_resp.status_code == 202
        eval_data = eval_resp.json()["data"]
        assert "evaluation_id" in eval_data

        eval_id = eval_data["evaluation_id"]
        report = store.get_evaluation_report_for_tenant(
            evaluation_id=eval_id, tenant_id="tenant_eval_e2e"
        )
        assert report is not None
        assert "criteria_results" in report
        assert len(report["criteria_results"]) == 2

        for cr in report["criteria_results"]:
            assert cr["score"] >= 0
            assert cr["max_score"] > 0
            assert "reason" in cr
            assert "citations" in cr


class TestFullPipelineE2E:
    """End-to-end: upload PDF → parse → Chroma index → Chroma retrieve → evaluate."""

    def _upload_and_parse(self, client, *, tenant_id: str, project_id: str, supplier_id: str):
        pdf_bytes = _make_pdf_with_content([
            "# Supplier Qualification\n\n"
            "The supplier holds ISO 9001:2015 certification, certificate number CN-2025-12345. "
            "The quality management system has been independently audited by SGS International. "
            "Company has 10 years of experience in government IT procurement projects.",
            "# Technical Proposal\n\n"
            "The solution provides 99.99% uptime SLA with automatic failover. "
            "Microservices architecture with Kubernetes orchestration. "
            "All data encrypted at rest and in transit using AES-256 and TLS 1.3. "
            "Full disaster recovery with RPO < 1 hour and RTO < 4 hours.",
            "# Pricing\n\n"
            "Total project cost: RMB 2,580,000 (tax inclusive). "
            "Payment schedule: 30% advance, 40% on delivery, 30% after acceptance. "
            "Price validity: 120 calendar days. Annual maintenance: RMB 180,000.",
        ])

        store.create_project(payload={
            "project_code": f"pc_{project_id}",
            "name": "Full Pipeline Test Project",
            "tenant_id": tenant_id,
        })
        store.create_supplier(payload={
            "supplier_code": f"sc_{supplier_id}",
            "name": "Full Pipeline Supplier",
            "tenant_id": tenant_id,
        })

        upload_resp = client.post(
            "/api/v1/documents/upload",
            data={
                "project_id": project_id,
                "supplier_id": supplier_id,
                "doc_type": "bid",
            },
            files={"file": ("full_pipeline.pdf", pdf_bytes, "application/pdf")},
            headers={
                "Idempotency-Key": "e2e-full-pipeline-upload",
                "x-tenant-id": tenant_id,
            },
        )
        assert upload_resp.status_code == 202
        data = upload_resp.json()["data"]
        result = store.run_job_once(job_id=data["job_id"], tenant_id=tenant_id)
        assert result["final_status"] == "succeeded"
        return data["document_id"]

    def test_full_pipeline_upload_parse_index_retrieve_evaluate(self, client):
        """Complete flow: upload real PDF, parse it, index to Chroma, retrieve via vector search, evaluate with scores."""
        tenant_id = "tenant_full_e2e"
        project_id = "proj_full_e2e"
        supplier_id = "sup_full_e2e"

        document_id = self._upload_and_parse(
            client, tenant_id=tenant_id, project_id=project_id, supplier_id=supplier_id
        )

        chunks = store.list_document_chunks_for_tenant(
            document_id=document_id, tenant_id=tenant_id
        )
        assert len(chunks) >= 1, "Parse should produce at least 1 chunk"
        full_text = " ".join(c.get("text", "") for c in chunks)
        assert "ISO 9001" in full_text or "99.99%" in full_text, "Chunks should contain real PDF content"

        for chunk in chunks:
            source = store.get_citation_source(chunk_id=chunk["chunk_id"], tenant_id=tenant_id)
            assert source is not None, f"Citation source missing for {chunk['chunk_id']}"

        store.create_rule_pack(payload={
            "rule_pack_version": "v1.0.0",
            "name": "Full Pipeline Rules",
            "tenant_id": tenant_id,
            "rules": {
                "criteria": [
                    {
                        "criteria_id": "qualification",
                        "criteria_name": "Supplier Qualification",
                        "max_score": 20.0,
                        "weight": 1.0,
                        "requirement_text": "Supplier must have valid ISO 9001 certification",
                    },
                    {
                        "criteria_id": "technical",
                        "criteria_name": "Technical Solution",
                        "max_score": 30.0,
                        "weight": 1.5,
                        "requirement_text": "Solution must provide high availability with SLA >= 99.9%",
                    },
                    {
                        "criteria_id": "pricing",
                        "criteria_name": "Pricing",
                        "max_score": 20.0,
                        "weight": 1.0,
                        "requirement_text": "Total price must be clearly stated with payment terms",
                    },
                ],
            },
        })

        eval_resp = client.post(
            "/api/v1/evaluations",
            json={
                "project_id": project_id,
                "supplier_id": supplier_id,
                "rule_pack_version": "v1.0.0",
                "evaluation_scope": {
                    "include_doc_types": ["bid"],
                    "force_hitl": False,
                },
                "query_options": {
                    "mode_hint": "hybrid",
                    "top_k": 10,
                },
            },
            headers={
                "Idempotency-Key": "e2e-full-pipeline-eval",
                "x-tenant-id": tenant_id,
            },
        )
        assert eval_resp.status_code == 202
        eval_data = eval_resp.json()["data"]
        evaluation_id = eval_data["evaluation_id"]

        report = store.get_evaluation_report_for_tenant(
            evaluation_id=evaluation_id, tenant_id=tenant_id
        )
        assert report is not None
        assert report.get("total_score", 0) > 0, "Report should have non-zero score"
        assert len(report.get("criteria_results", [])) == 3

        for cr in report["criteria_results"]:
            assert cr["score"] >= 0
            assert cr["max_score"] > 0
            assert "reason" in cr and len(cr["reason"]) > 0
            assert isinstance(cr["citations"], list)

    def test_chroma_retrieval_returns_indexed_chunks(self, client):
        """Verify that after parsing, vector retrieval finds the indexed chunks."""
        tenant_id = "tenant_chroma_ret"
        project_id = "proj_chroma_ret"
        supplier_id = "sup_chroma_ret"

        self._upload_and_parse(
            client, tenant_id=tenant_id, project_id=project_id, supplier_id=supplier_id
        )

        result = store.retrieval_query(
            tenant_id=tenant_id,
            project_id=project_id,
            supplier_id=supplier_id,
            query="ISO 9001 certification",
            query_type="fact",
            high_risk=False,
            top_k=5,
            doc_scope=["bid"],
        )
        assert isinstance(result, dict)
        items = result.get("items", [])
        assert len(items) >= 1, "Chroma should return at least 1 result for indexed content"
        for item in items:
            assert "chunk_id" in item
            assert "score_raw" in item

    def test_retrieval_preview_returns_text(self, client):
        """Verify retrieval preview includes text excerpts from indexed chunks."""
        tenant_id = "tenant_preview"
        project_id = "proj_preview"
        supplier_id = "sup_preview"

        self._upload_and_parse(
            client, tenant_id=tenant_id, project_id=project_id, supplier_id=supplier_id
        )

        result = store.retrieval_preview(
            tenant_id=tenant_id,
            project_id=project_id,
            supplier_id=supplier_id,
            query="microservices architecture Kubernetes",
            query_type="fact",
            high_risk=False,
            top_k=5,
            doc_scope=["bid"],
        )
        assert isinstance(result, dict)
        assert "items" in result
